"""Document ingestion pipeline: load → split → embed → sync into Supabase.

Markdown dosyaları header-bazlı bölünür (her bölüm anlamlı bir bütün, embedlenen
metne "Projeler > Business Data Finder" gibi breadcrumb prefix'i eklenir).
Diğer formatlar (pdf/docx/txt) karakter-bazlı generic splitter'dan geçer.

Senkronizasyon idempotenttir: her kaynak için önce delete-by-source, sonra insert.
Aynı dosyayı tekrar ingest etmek mükerrer satır üretmez.

CLI: python -m app.ingest <file_or_dir> [--source <label>] [--wipe]
"""
from __future__ import annotations

import argparse
import asyncio
import hashlib
import os
from collections.abc import Iterable
from datetime import UTC, datetime
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

from .config import Settings, get_settings
from .deps import embeddings, supabase_client
from .retry import external_retry

# Bilgi tabaninin kok dizini. /admin/ingest-path yalnizca bunun altini kabul eder
# (bkz. routes/admin.py::_resolve_data_path); ingest_path'in kendisi kisitsiz kalir
# cunku /admin/ingest yuklenen dosyayi tempfile yoluyla, CLI ise operatorun verdigi
# herhangi bir yolu mesru sekilde kullanir.
DATA_ROOT = Path(__file__).resolve().parents[1] / "data"

_MD_EXTS = {".md", ".markdown"}
_SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".html", ".htm"} | _MD_EXTS

_HEADERS_TO_SPLIT_ON = [("#", "h1"), ("##", "h2"), ("###", "h3")]


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def _read_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext in {".txt", ".html", ".htm"} | _MD_EXTS:
        return _read_text(path)
    raise ValueError(f"Unsupported file type: {ext}")


def _iter_files(root: Path) -> Iterable[Path]:
    if root.is_file():
        yield root
        return
    for p in sorted(root.rglob("*")):
        if p.name.lower() == "readme.md":  # yazım kuralları dosyası, bilgi tabanına girmez
            continue
        if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTS:
            yield p


def _finalize(chunks: list[tuple[str, list[str]]], source: str) -> list[Document]:
    """(text, header_path) çiftlerini ortak metadata şemasıyla Document'a çevirir."""
    now = datetime.now(UTC).isoformat()
    docs: list[Document] = []
    for i, (text, headers) in enumerate(chunks):
        docs.append(Document(
            page_content=text,
            metadata={
                "source": source,
                "headers": headers,
                "chunk_index": i,
                "content_hash": hashlib.sha256(text.encode("utf-8")).hexdigest()[:16],
                "ingested_at": now,
            },
        ))
    return docs


def chunk_markdown(text: str, source: str, s: Settings) -> list[Document]:
    """Header-bazlı bölme; CHUNK_SIZE'ı aşan bölümler ikincil splitter'dan geçer.

    Embedlenen metnin başına header breadcrumb'ı eklenir — kısa sorgularda
    bölüm adının (örn. proje adı) recall'unu garanti eden asıl mekanizma budur.
    """
    md_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=_HEADERS_TO_SPLIT_ON, strip_headers=False
    )
    fallback = RecursiveCharacterTextSplitter(
        chunk_size=s.CHUNK_SIZE, chunk_overlap=s.CHUNK_OVERLAP
    )
    chunks: list[tuple[str, list[str]]] = []
    for section in md_splitter.split_text(text):
        headers = [section.metadata[key] for _, key in _HEADERS_TO_SPLIT_ON
                   if key in section.metadata]
        breadcrumb = " > ".join(headers)
        pieces = ([section.page_content] if len(section.page_content) <= s.CHUNK_SIZE
                  else fallback.split_text(section.page_content))
        for piece in pieces:
            body = piece if not breadcrumb else f"{breadcrumb}\n\n{piece}"
            chunks.append((body, headers))
    return _finalize(chunks, source)


def chunk_generic(text: str, source: str, s: Settings) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=s.CHUNK_SIZE, chunk_overlap=s.CHUNK_OVERLAP
    )
    return _finalize([(c, []) for c in splitter.split_text(text)], source)


def chunk_file(path: Path, source: str, s: Settings) -> list[Document]:
    text = _load(path)
    if not text.strip():
        return []
    if path.suffix.lower() in _MD_EXTS:
        return chunk_markdown(text, source, s)
    return chunk_generic(text, source, s)


@external_retry
def _embed_texts(texts: list[str]) -> list[list[float]]:
    return embeddings().embed_documents(texts)


@external_retry
def _delete_source(source: str) -> None:
    s = get_settings()
    supabase_client().table(s.SUPABASE_TABLE).delete().eq("metadata->>source", source).execute()


@external_retry
def _insert_rows(rows: list[dict]) -> None:
    s = get_settings()
    supabase_client().table(s.SUPABASE_TABLE).insert(rows).execute()


@external_retry
def _wipe_table() -> None:
    s = get_settings()
    # PostgREST filtresiz delete kabul etmez; tüm id'leri kapsayan koşul kullanılır.
    supabase_client().table(s.SUPABASE_TABLE).delete().gte("id", 0).execute()


_INSERT_BATCH = 100


async def sync_source(source: str, docs: list[Document]) -> int:
    """Kaynağın eski chunk'larını silip yenilerini yazar (idempotent)."""
    vectors = await asyncio.to_thread(_embed_texts, [d.page_content for d in docs])
    # strict=True: embedding sayisi chunk sayisiyla uyusmazsa sessizce kirpmak yerine
    # patlasin — kisa donen bir embedding cevabi korpusun kuyrugunu yazmadan gecerdi.
    rows = [
        {"content": d.page_content, "metadata": d.metadata, "embedding": v}
        for d, v in zip(docs, vectors, strict=True)
    ]
    await asyncio.to_thread(_delete_source, source)
    for i in range(0, len(rows), _INSERT_BATCH):
        await asyncio.to_thread(_insert_rows, rows[i:i + _INSERT_BATCH])
    return len(rows)


async def ingest_path(target: str, source_label: str | None = None, wipe: bool = False) -> int:
    root = Path(target)
    if not root.exists():
        raise FileNotFoundError(target)

    s = get_settings()
    by_source: dict[str, list[Document]] = {}
    for f in _iter_files(root):
        source = source_label or os.path.relpath(
            str(f), start=str(root.parent if root.is_file() else root)
        )
        docs = chunk_file(f, source, s)
        if docs:
            by_source.setdefault(source, []).extend(docs)

    if wipe:
        await asyncio.to_thread(_wipe_table)

    total = 0
    for source, docs in by_source.items():
        total += await sync_source(source, docs)
    return total


def _main() -> None:
    parser = argparse.ArgumentParser(description="Ingest documents into Supabase vector store")
    parser.add_argument("path", help="File or directory")
    parser.add_argument("--source", help="Optional source label", default=None)
    parser.add_argument("--wipe", action="store_true",
                        help="Önce TÜM tabloyu temizle (eski kaynakların tek seferlik temizliği)")
    args = parser.parse_args()
    n = asyncio.run(ingest_path(args.path, args.source, wipe=args.wipe))
    print(f"Ingested {n} chunks from {args.path}")


if __name__ == "__main__":
    _main()
